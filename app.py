# app.py
import os
import time
import math
import re
import logging
import subprocess
import shutil
import sys
from pathlib import Path
from shutil import copyfile
from datetime import datetime
from functools import wraps
from io import BytesIO

from flask import (
    Flask, render_template, request, send_file, jsonify, abort,
    url_for, redirect, session, flash
)
from werkzeug.utils import secure_filename

from docx import Document
from num2words import num2words

# TOTP & QR
import pyotp
import qrcode

# Optional docx2pdf (we import dynamically inside try to avoid startup failure)
# Optional Windows COM init/uninit
try:
    import pythoncom
    import pywintypes
except Exception:
    pythoncom = None
    pywintypes = None

# ---------- Config ----------
# If APP_SECRET env var is NOT provided, we generate a random secret at startup.
# That will invalidate cookies after every restart (forces TOTP again).
APP_SECRET = os.environ.get("APP_SECRET")
if not APP_SECRET:
    # random secret each start -> forces login after restart
    APP_SECRET = os.urandom(32).hex()

TOTP_SECRET_FILE = Path("auth_totp_secret.txt")
UPLOAD_TEMPLATE = "uploads/template.docx"
OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

app = Flask(__name__)
app.secret_key = APP_SECRET

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ---------- Auth helpers (passcode-only TOTP) ----------
def get_totp_secret():
    if TOTP_SECRET_FILE.exists():
        return TOTP_SECRET_FILE.read_text().strip()
    return None

def save_totp_secret(secret):
    TOTP_SECRET_FILE.write_text(secret)

def generate_totp_uri(secret, issuer_name="GBUILDCON", account_name="GB-Admin"):
    return pyotp.totp.TOTP(secret).provisioning_uri(name=account_name, issuer_name=issuer_name)

def verify_totp_code(secret, code):
    try:
        totp = pyotp.TOTP(secret)
        return totp.verify(code, valid_window=1)  # allow ±1 step skew
    except Exception:
        return False

def login_required(f):
    @wraps(f)
    def wrapped(*args, **kwargs):
        if not session.get("authenticated"):
            return redirect(url_for("login", next=request.path))
        return f(*args, **kwargs)
    return wrapped

# ---------- Number to words (Indian) ----------
def number_to_words_indian(amount):
    amount_str = "{:.2f}".format(float(amount))
    rupees_str, paise_str = amount_str.split(".")
    rupees = int(rupees_str)
    paise = int(paise_str)
    rupees_words = num2words(rupees, lang='en_IN').replace(",", "").title()
    if paise > 0:
        paise_words = num2words(paise, lang='en_IN').replace(",", "").title()
        return f"{rupees_words} Rupees And {paise_words} Paise Only"
    else:
        return f"{rupees_words} Rupees Only"

# ---------- Placeholder replacement & doc helpers ----------
PLACEHOLDER_RE = re.compile(r'(\{\{[^}]+\}\})')

BOLD_KEYS = {
    "{{base_amo}}","{{cgst_c}}","{{sgst_c}}","{{tax_t}}","{{sub_total_words}}","{{tax_to_words}}"
}
BOLD_PARAGRAPH_KEYS = {"{{sub_total_words}}","{{tax_to_words}}","{{sub_total}}"}

def replace_placeholders_in_paragraph(paragraph, replacements):
    full_text = ''.join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return
    paragraph_should_be_bold = any(k in full_text for k in BOLD_PARAGRAPH_KEYS)
    segments=[]
    idx=0
    found=False
    for m in PLACEHOLDER_RE.finditer(full_text):
        found=True
        start,end=m.span()
        placeholder=m.group(1)
        if start>idx:
            segments.append((full_text[idx:start], paragraph_should_be_bold))
        replacement_text = replacements.get(placeholder, placeholder)
        is_bold = paragraph_should_be_bold or (placeholder in BOLD_KEYS)
        segments.append((replacement_text, is_bold))
        idx=end
    if idx < len(full_text):
        segments.append((full_text[idx:], paragraph_should_be_bold))
    if not found:
        return
    for r in paragraph.runs:
        r.text=""
    for text, bold_flag in segments:
        if not text:
            continue
        r = paragraph.add_run(text)
        if bold_flag:
            r.bold=True

def replace_placeholders_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            replace_placeholders_in_block(cell, replacements)

def replace_placeholders_in_block(block, replacements):
    if hasattr(block, "paragraphs"):
        for p in list(block.paragraphs):
            replace_placeholders_in_paragraph(p, replacements)
    if hasattr(block, "tables"):
        for t in block.tables:
            replace_placeholders_in_table(t, replacements)

def replace_placeholders_in_doc(doc, replacements):
    replace_placeholders_in_block(doc, replacements)
    for section in doc.sections:
        try:
            replace_placeholders_in_block(section.header, replacements)
        except Exception:
            pass
        try:
            replace_placeholders_in_block(section.footer, replacements)
        except Exception:
            pass

def is_paragraph_empty(paragraph):
    if paragraph.text and paragraph.text.strip():
        return False
    for run in paragraph.runs:
        if run.text and run.text.strip():
            return False
    return True

def remove_trailing_empty_paragraphs_from_block(block):
    if not hasattr(block, "paragraphs"):
        return
    while block.paragraphs:
        last = block.paragraphs[-1]
        if is_paragraph_empty(last):
            p_el = last._element
            parent = p_el.getparent()
            if parent is not None:
                parent.remove(p_el)
            else:
                break
        else:
            break

def remove_trailing_empty_paragraphs(doc):
    remove_trailing_empty_paragraphs_from_block(doc)
    for section in doc.sections:
        try:
            remove_trailing_empty_paragraphs_from_block(section.header)
        except Exception:
            pass
        try:
            remove_trailing_empty_paragraphs_from_block(section.footer)
        except Exception:
            pass

# ---------- Invoice list & suffix (only PDFs) ----------
def list_existing_invoices():
    invoices=[]
    for p in sorted(Path(OUTPUT_DIR).glob("*.pdf"), reverse=True):
        safe_name=p.name
        display_name=safe_name.replace("_"," ")
        m=re.search(r"(GB-\d{4}-\d+)", safe_name)
        invoice_no=m.group(1) if m else ""
        try:
            url=url_for("download_file", filename=safe_name)
        except RuntimeError:
            url=f"/download/{safe_name}"
        invoices.append({"filename":safe_name,"display_name":display_name,"invoice_no":invoice_no,"url":url})
    return invoices

def get_next_suffix_for_month(month_year):
    max_suffix=0
    pattern=re.compile(rf"GB-{re.escape(month_year)}-(\d+)")
    for path in Path(OUTPUT_DIR).glob("*.pdf"):
        name=path.stem
        m=pattern.search(name)
        if m:
            try:
                val=int(m.group(1))
                if val>max_suffix:
                    max_suffix=val
            except Exception:
                pass
    return f"{max_suffix+1:02d}"

# ---------- Auth routes ----------
@app.route("/setup-2fa", methods=["GET","POST"])
def setup_2fa():
    existing = get_totp_secret()
    if request.method == "GET":
        if existing:
            return render_template("setup_2fa.html", already_setup=True)
        secret = pyotp.random_base32()
        session["setup_tmp_secret"] = secret
        return render_template("setup_2fa.html", already_setup=False)
    secret = session.get("setup_tmp_secret")
    if not secret:
        flash("No setup in progress. Reload page.", "danger")
        return redirect(url_for("setup_2fa"))
    code = (request.form.get("code") or "").strip()
    if verify_totp_code(secret, code):
        save_totp_secret(secret)
        session.pop("setup_tmp_secret", None)
        flash("TOTP configured. Now use the code on the login page.", "success")
        return redirect(url_for("login"))
    else:
        flash("Invalid code. Try again.", "danger")
        return redirect(url_for("setup_2fa"))

@app.route("/setup-qr.png")
def setup_qr_png():
    secret = session.get("setup_tmp_secret")
    if not secret:
        abort(404)
    uri = generate_totp_uri(secret)
    img = qrcode.make(uri)
    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return send_file(buf, mimetype="image/png")

@app.route("/login", methods=["GET","POST"])
def login():
    next_url = request.args.get("next") or url_for("index")
    totp_secret = get_totp_secret()
    if request.method == "GET":
        return render_template("login.html", totp_enabled=bool(totp_secret), next=next_url)
    code = (request.form.get("code") or "").strip()
    secret = get_totp_secret()
    if not secret:
        flash("TOTP not configured. Visit Setup 2FA.", "danger")
        return redirect(url_for("setup_2fa"))
    if not code:
        flash("Enter the 6-digit code", "danger")
        return redirect(url_for("login"))
    if not verify_totp_code(secret, code):
        flash("Invalid passcode", "danger")
        return redirect(url_for("login"))
    session["authenticated"] = True
    flash("Authenticated", "success")
    return redirect(next_url)

@app.route("/logout")
def logout():
    session.clear()
    flash("Logged out", "info")
    return redirect(url_for("login"))

# ---------- Protected invoice routes ----------
@app.route('/')
@login_required
def index():
    today = datetime.now().strftime("%Y-%m-%d")
    month_year = datetime.now().strftime("%m%y")
    next_suffix = get_next_suffix_for_month(month_year)
    invoices = list_existing_invoices()
    return render_template("index.html",
                           default_date=today,
                           suggested_suffix=next_suffix,
                           invoices=invoices)

@app.route('/next-suffix')
@login_required
def next_suffix():
    date_value = request.args.get("date")
    if not date_value:
        return jsonify({"error":"date is required"}),400
    try:
        month_year = datetime.strptime(date_value, "%Y-%m-%d").strftime("%m%y")
    except Exception:
        return jsonify({"error":"invalid date"}),400
    next_suffix = get_next_suffix_for_month(month_year)
    return jsonify({"next_suffix": next_suffix})

@app.route('/generate', methods=['POST'])
@login_required
def generate_invoice():
    # Read amount and compute
    try:
        y_value = float(request.form['y_value'])
    except Exception:
        return "Invalid y_value", 400
    sub_total = math.ceil(y_value * 1.02)

    date_value = request.form.get('date', datetime.now().strftime("%Y-%m-%d"))
    invoice_suffix = request.form.get('invoice_suffix','').strip()
    payment_type = request.form.get('payment_type','')
    hsn_code = request.form.get('hsn_code','')

    if not invoice_suffix.isdigit():
        return "Invoice suffix must be numeric", 400

    date_formatted = datetime.strptime(date_value, "%Y-%m-%d").strftime("%d/%m/%Y")
    month_year = datetime.strptime(date_value, "%Y-%m-%d").strftime("%m%y")
    invoice_no = f"GB-{month_year}-{invoice_suffix}"

    base_amount = sub_total / 1.18
    cgst = base_amount * 0.09
    sgst = base_amount * 0.09
    sub_total_2 = base_amount + cgst + sgst
    tax_total = cgst + sgst

    sub_total_words = number_to_words_indian(sub_total)
    tax_total_words = number_to_words_indian(tax_total)

    # Load template
    try:
        doc = Document(UPLOAD_TEMPLATE)
    except Exception as e:
        logger.exception("Failed to open template: %s", e)
        return f"Failed to open template: {e}", 500

    replacements = {
        "{{BASE_TOTAL}}": f"{base_amount:.2f}",
        "{{BASE_TO}}": f"{base_amount:.2f}",
        "{{SUB_TOTAL}}": f"{sub_total:.2f}",
        "{{TAX_TO_C}}": f"{tax_total:.2f}",
        "{{SGST_C}}": f"{sgst:.2f}",
        "{{CGST_C}}": f"{cgst:.2f}",
        "{{base_amount}}": f"{base_amount:.2f}",
        "{{base_amo}}": f"{base_amount:.2f}",
        "{{cgst}}": f"{cgst:.2f}",
        "{{sgst}}": f"{sgst:.2f}",
        "{{tax_to}}": f"{tax_total:.2f}",
        "{{cgst_c}}": f"{cgst:.2f}",
        "{{sgst_c}}": f"{sgst:.2f}",
        "{{tax_t}}": f"{tax_total:.2f}",
        "{{sub_total}}": f"{sub_total_2:.2f}",
        "{{sub_total_words}}": sub_total_words,
        "{{tax_to_words}}": tax_total_words,
        "{{date}}": date_formatted,
        "{{invoice_no}}": invoice_no,
        "{{payment_type}}": payment_type,
        "{{hsn}}": hsn_code,
        "{{hsn_1}}": hsn_code,
        "{{hsn_n}}": hsn_code,
    }

    replace_placeholders_in_doc(doc, replacements)
    remove_trailing_empty_paragraphs(doc)

    # Save temporary DOCX to convert
    timestamp = int(time.time())
    tmp_docx = os.path.join(OUTPUT_DIR, f"tmp_generated_{timestamp}.docx")
    try:
        doc.save(tmp_docx)
    except Exception as e:
        logger.exception("Failed to save DOCX: %s", e)
        return f"Failed to save docx: {e}", 500

    # Prepare PDF name (friendly) and safe filename
    pdf_filename = f"G BUILDCON Invoice - {invoice_no}.pdf"
    safe_pdf_filename = secure_filename(pdf_filename)
    output_pdf = os.path.join(OUTPUT_DIR, safe_pdf_filename)

    # ---------- Conversion: try docx2pdf (Word/COM) first, with COM init, then LibreOffice fallback ----------
    conversion_error = None
    converted_ok = False

    # Try docx2pdf (Word/COM) first — but ensure COM is initialized on Windows
    try:
        if pythoncom is not None:
            try:
                pythoncom.CoInitialize()
            except Exception as e:
                logger.warning("pythoncom.CoInitialize() warning: %s", e)

        try:
            # Import docx2pdf dynamically
            from docx2pdf import convert as docx2pdf_convert
            try:
                docx2pdf_convert(tmp_docx, output_pdf)
                logger.info("Converted DOCX to PDF via docx2pdf: %s", output_pdf)
                converted_ok = True
            except Exception as e:
                conversion_error = f"docx2pdf conversion failed: {e}"
                logger.exception("docx2pdf conversion failed: %s", e)
        except Exception as e:
            # docx2pdf not available or import failed
            conversion_error = f"docx2pdf import failed: {e}"
            logger.exception("docx2pdf import/usage failed: %s", e)
    finally:
        if pythoncom is not None:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

    # If docx2pdf didn't work, try LibreOffice (soffice) headless conversion as a fallback
    if not converted_ok:
        try:
            # Find soffice/libreoffice binary
            soffice_cmd = None
            for cmd in ("soffice", "libreoffice", "lowriter"):
                if shutil.which(cmd):
                    soffice_cmd = cmd
                    break

            if not soffice_cmd:
                raise FileNotFoundError("LibreOffice 'soffice' not found on PATH")

            # Run conversion
            run = subprocess.run(
                [soffice_cmd, "--headless", "--convert-to", "pdf", tmp_docx, "--outdir", OUTPUT_DIR],
                check=True,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )

            # LibreOffice output file name: same base name, .pdf extension, located in OUTPUT_DIR
            tmp_pdf_name = Path(tmp_docx).with_suffix(".pdf").name
            tmp_pdf_path = os.path.join(OUTPUT_DIR, tmp_pdf_name)

            if not os.path.exists(tmp_pdf_path):
                raise FileNotFoundError(f"LibreOffice conversion claimed success but {tmp_pdf_path} not found. stdout: {run.stdout} stderr: {run.stderr}")

            # Move/rename to desired output_pdf name if needed
            if os.path.abspath(tmp_pdf_path) != os.path.abspath(output_pdf):
                try:
                    if os.path.exists(output_pdf):
                        os.remove(output_pdf)
                    shutil.move(tmp_pdf_path, output_pdf)
                except Exception as e:
                    raise RuntimeError(f"Failed to move LibreOffice PDF: {e}")

            logger.info("Converted DOCX to PDF via LibreOffice: %s", output_pdf)
            converted_ok = True

        except Exception as e:
            conversion_error = (conversion_error or "") + f" | LibreOffice conversion failed: {e}"
            logger.exception("LibreOffice conversion failed: %s", e)

    # Remove temporary docx in all cases (we keep output folder PDF-only)
    try:
        if os.path.exists(tmp_docx):
            os.remove(tmp_docx)
    except Exception:
        pass

    if not converted_ok:
        msg = (
            "Failed to convert to PDF. "
            "On Windows ensure MS Word is installed and accessible; on Linux install LibreOffice. "
            "Conversion details: " + str(conversion_error)
        )
        return (msg, 500)

    # Send the PDF as download
    return send_file(
        output_pdf,
        as_attachment=True,
        download_name=pdf_filename,
        mimetype="application/pdf"
    )

@app.route('/download/<path:filename>')
@login_required
def download_file(filename):
    # Validate and ensure only PDFs listed
    if "/" in filename or "\\" in filename or ".." in filename:
        abort(400)
    full_path = os.path.join(OUTPUT_DIR, filename)
    try:
        resolved = Path(full_path).resolve()
        if not str(resolved).startswith(str(Path(OUTPUT_DIR).resolve())):
            abort(400)
    except Exception:
        abort(400)
    if not os.path.exists(full_path):
        abort(404)
    # Only allow PDFs for download (safe)
    if not filename.lower().endswith(".pdf"):
        abort(400)
    friendly_name = filename.replace("_", " ")
    return send_file(full_path, as_attachment=True, download_name=friendly_name, mimetype="application/pdf")

if __name__ == "__main__":
    app.run(debug=True)
